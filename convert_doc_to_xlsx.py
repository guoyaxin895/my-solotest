import os
from docx import Document
from openpyxl import Workbook

def convert_doc_to_xlsx(doc_path, xlsx_path):
    if not os.path.exists(doc_path):
        print(f"错误：文件 {doc_path} 不存在")
        return False

    doc = Document(doc_path)
    wb = Workbook()
    ws = wb.active
    ws.title = "Content"

    row = 1
    for para in doc.paragraphs:
        if para.text.strip():
            ws.cell(row=row, column=1, value=para.text)
            row += 1

    for table in doc.tables:
        for table_row in table.rows:
            for cell in table_row.cells:
                if cell.text.strip():
                    ws.cell(row=row, column=1, value=cell.text)
                    row += 1

    wb.save(xlsx_path)
    print(f"转换成功：{xlsx_path}")
    return True

if __name__ == "__main__":
    script_dir = os.path.dirname(os.path.abspath(__file__))
    doc_file = os.path.join(script_dir, "2026-05-08.doc")
    xlsx_file = os.path.join(script_dir, "2026-05-08.xlsx")

    success = convert_doc_to_xlsx(doc_file, xlsx_file)
    if not success:
        exit(1)
