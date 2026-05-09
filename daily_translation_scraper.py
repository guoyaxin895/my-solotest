import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from datetime import datetime
import os

def fetch_daily_translation():
    base_url = "https://fanyi.kkabc.com"
    list_url = f"{base_url}/mrfy/"

    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
    }

    response = requests.get(list_url, headers=headers, timeout=10)
    response.encoding = 'utf-8'

    soup = BeautifulSoup(response.text, 'html.parser')

    all_links = soup.find_all('a')
    mrfy_links = [a for a in all_links if '/mrfy/show-' in a.get('href', '')]

    if not mrfy_links:
        print("未找到每日翻译文章链接")
        return

    first_article_link = mrfy_links[0]

    article_url = base_url + first_article_link['href']
    article_response = requests.get(article_url, headers=headers, timeout=10)
    article_response.encoding = 'utf-8'

    article_soup = BeautifulSoup(article_response.text, 'html.parser')

    title = article_soup.find('h2').get_text(strip=True)

    content_div = article_soup.find('div', class_='show-text')
    content_text = content_div.get_text(separator='\n', strip=True) if content_div else ""

    doc = Document()
    doc.add_heading(title, level=1)

    paragraphs = content_text.split('\n')
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para.strip())

    today = datetime.now().strftime('%Y-%m-%d')
    filename = f"{today}.doc"
    doc.save(filename)
    print(f"文章已保存为: {filename}")

if __name__ == "__main__":
    fetch_daily_translation()
